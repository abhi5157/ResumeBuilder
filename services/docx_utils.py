"""
Utility functions for creating and styling DOCX documents.
Provides consistent formatting for resume generation.
"""

from typing import Optional, List, Tuple, Dict
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_page_margins(doc: Document, 
                     top: float = 0.5, 
                     bottom: float = 0.5, 
                     left: float = 0.5, 
                     right: float = 0.5):
    """
    Set page margins in inches.
    
    Args:
        doc: Document object
        top: Top margin in inches
        bottom: Bottom margin in inches
        left: Left margin in inches
        right: Right margin in inches
    """
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


def add_heading_with_line(doc: Document, text: str, font_size: int = 11) -> None:
    """
    Add a section heading with a horizontal line above it.
    
    Args:
        doc: Document object
        text: Heading text
        font_size: Font size in points
    """
    # Add horizontal line first
    line_para = doc.add_paragraph()
    line_para.paragraph_format.space_before = Pt(0)
    line_para.paragraph_format.space_after = Pt(0)
    
    # Add top border to create horizontal line
    pPr = line_para._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '0')
    top.set(qn('w:color'), '000000')
    pBdr.append(top)
    pPr.append(pBdr)
    
    # Add heading text below the line
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text.upper())
    run.bold = False
    run.font.size = Pt(font_size)
    run.font.name = 'Calibri Light'
    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Set spacing
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0


def add_contact_header(doc: Document, 
                       name: str, 
                       email: str, 
                       phone: str, 
                       location: str,
                       linkedin: Optional[str] = None,
                       clearance: Optional[str] = None) -> None:
    """
    Add contact header at the top of the resume.
    
    Args:
        doc: Document object
        name: Full name
        email: Email address
        phone: Phone number
        location: City, State
        linkedin: LinkedIn URL
        clearance: Security clearance level
    """
    # Name - centered, large, bold
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(name.upper())
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_run.font.name = 'Calibri Light'
    name_para.paragraph_format.space_after = Pt(2)
    
    # Contact info - centered, smaller
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    contact_parts = [location, phone, email]
    if linkedin:
        contact_parts.append(linkedin)
    if clearance and clearance != 'None':
        contact_parts.append(f"Clearance: {clearance}")
    
    contact_text = ' | '.join(contact_parts)
    contact_run = contact_para.add_run(contact_text)
    contact_run.font.size = Pt(10)
    contact_run.font.name = 'Calibri Light'
    contact_para.paragraph_format.space_after = Pt(2)


def add_branch_title(doc: Document, branch_title: str) -> None:
    """
    Add military branch title (e.g., "UNITED STATES NAVY - SURFACE WARFARE OFFICER").
    
    Args:
        doc: Document object
        branch_title: Branch and MOS title
    """
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(branch_title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Calibri Light'
    para.paragraph_format.space_after = Pt(2)


def add_summary_section(doc: Document, summary: str) -> None:
    """
    Add professional summary section.
    
    Args:
        doc: Document object
        summary: Summary text
    """
    add_heading_with_line(doc, "SUMMARY")
    
    para = doc.add_paragraph(summary)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(2)
    
    for run in para.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri Light'


def add_experience_entry(doc: Document,
                        title: str,
                        subtitle: str,
                        location: str,
                        date_range: str,
                        bullets: List[str]) -> None:
    """
    Add a single work experience entry.
    
    Args:
        doc: Document object
        title: Job title
        subtitle: Company/Organization
        location: Location
        date_range: Date range (e.g., "January 2020 - Present")
        bullets: List of achievement bullets
    """
    # Create table for title row (title on left, date on right)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Calculate full width (page width - margins)
    # Page width is 8.5", left margin 0.7", right margin 0.7" = 7.1" available
    full_width = Inches(7.1)
    
    # Set column widths with 75/25 ratio
    table.columns[0].width = Inches(5.325)  # 75% of 7.1"
    table.columns[1].width = Inches(1.775)  # 25% of 7.1"
    
    # Add vertical border between columns only
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = tcPr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement('w:tcBorders')
                tcPr.append(borders)
            
            # Add vertical divider between cells
            if idx == 0:
                # Right border on left cell
                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'single')
                right.set(qn('w:sz'), '6')
                right.set(qn('w:space'), '0')
                right.set(qn('w:color'), '000000')
                borders.append(right)
            else:
                # Left border on right cell
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single')
                left.set(qn('w:sz'), '6')
                left.set(qn('w:space'), '0')
                left.set(qn('w:color'), '000000')
                borders.append(left)
            
            # Remove other borders
            for edge in ['top', 'bottom']:
                edge_elem = OxmlElement(f'w:{edge}')
                edge_elem.set(qn('w:val'), 'none')
                edge_elem.set(qn('w:sz'), '0')
                borders.append(edge_elem)
    
    # Title cell (left)
    title_cell = table.rows[0].cells[0]
    title_para = title_cell.paragraphs[0]
    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_run.font.name = 'Calibri Light'
    title_para.paragraph_format.space_after = Pt(0)
    
    # Add organization and location on the same line
    title_para.add_run(', ')
    org_run = title_para.add_run(subtitle)
    org_run.font.italic = True
    org_run.font.size = Pt(10)
    org_run.font.name = 'Calibri Light'
    org_run.bold = False
    
    if location and location != subtitle:
        title_para.add_run(', ')
        loc_run = title_para.add_run(location)
        loc_run.font.italic = True
        loc_run.font.size = Pt(10)
        loc_run.font.name = 'Calibri Light'
        loc_run.bold = False
    
    # Add bullets inside the left cell
    for bullet in bullets:
        bullet_para = title_cell.add_paragraph(f'• {bullet}')
        bullet_para.paragraph_format.left_indent = Inches(0.15)
        bullet_para.paragraph_format.first_line_indent = Inches(-0.15)
        bullet_para.paragraph_format.space_before = Pt(0)
        bullet_para.paragraph_format.space_after = Pt(0)
        bullet_para.paragraph_format.line_spacing = 1.0
        
        for run in bullet_para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri Light'
    
    # Date cell (right)
    date_cell = table.rows[0].cells[1]
    date_para = date_cell.paragraphs[0]
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(date_range)
    date_run.font.size = Pt(10)
    date_run.font.name = 'Calibri Light'
    date_run.italic = False
    date_para.paragraph_format.space_after = Pt(0)
    
    # Add spacing after table
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_after = Pt(2)


def add_experience_section_table(doc: Document, experiences: List[Dict]) -> None:
    """
    Add all work experience entries in a single table with shared outer border.
    
    Args:
        doc: Document object
        experiences: List of experience dictionaries with keys: title, subtitle, location, date_range, bullets
    """
    if not experiences:
        return
    
    # Create table with number of rows = number of experiences
    table = doc.add_table(rows=len(experiences), cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Calculate full width (page width - margins)
    full_width = Inches(7.1)
    
    # Set column widths with 75/25 ratio
    table.columns[0].width = Inches(5.325)  # 75% of 7.1"
    table.columns[1].width = Inches(1.775)  # 25% of 7.1"
    
    # Add borders to the entire table
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Add table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Fill in each row with experience data
    for idx, exp in enumerate(experiences):
        row = table.rows[idx]
        
        # Left cell (content)
        left_cell = row.cells[0]
        left_para = left_cell.paragraphs[0]
        
        # Title
        title_run = left_para.add_run(exp['title'])
        title_run.bold = True
        title_run.font.size = Pt(11)
        title_run.font.name = 'Calibri Light'
        left_para.paragraph_format.space_after = Pt(0)
        
        # Organization and location
        left_para.add_run(', ')
        org_run = left_para.add_run(exp['subtitle'])
        org_run.font.italic = True
        org_run.font.size = Pt(10)
        org_run.font.name = 'Calibri Light'
        org_run.bold = False
        
        if exp['location'] and exp['location'] != exp['subtitle']:
            left_para.add_run(', ')
            loc_run = left_para.add_run(exp['location'])
            loc_run.font.italic = True
            loc_run.font.size = Pt(10)
            loc_run.font.name = 'Calibri Light'
            loc_run.bold = False
        
        # Add bullets
        for bullet in exp['bullets']:
            bullet_para = left_cell.add_paragraph(f'• {bullet}')
            bullet_para.paragraph_format.left_indent = Inches(0.15)
            bullet_para.paragraph_format.first_line_indent = Inches(-0.15)
            bullet_para.paragraph_format.space_before = Pt(0)
            bullet_para.paragraph_format.space_after = Pt(0)
            bullet_para.paragraph_format.line_spacing = 1.0
            
            for run in bullet_para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri Light'
        
        # Right cell (date)
        right_cell = row.cells[1]
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_run = right_para.add_run(exp['date_range'])
        date_run.font.size = Pt(10)
        date_run.font.name = 'Calibri Light'
        date_run.italic = False
        right_para.paragraph_format.space_after = Pt(0)
    
    # Add spacing after table
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_after = Pt(4)


def add_education_certification_combined_table(doc: Document, education_list: List[Dict], certifications: List[Dict]) -> None:
    """
    Add education and certification entries in a single combined table with shared outer border.
    Education entries have content on left, graduation on right; certifications are two per row side by side in left column.
    
    Args:
        doc: Document object
        education_list: List of education dictionaries
        certifications: List of certification dictionaries
    """
    if not education_list and not certifications:
        return
    
    # Calculate total rows needed
    education_rows = len(education_list) if education_list else 0
    cert_rows = (len(certifications) + 1) // 2 if certifications else 0  # Two certifications per row in left column
    total_rows = education_rows + cert_rows
    
    if total_rows == 0:
        return
    
    # Create table with total rows and 2 columns
    table = doc.add_table(rows=total_rows, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Calculate full width (page width - margins)
    full_width = Inches(7.1)
    
    # Set column widths with 75/25 ratio like experience
    table.columns[0].width = Inches(5.325)  # 75% of 7.1"
    table.columns[1].width = Inches(1.775)  # 25% of 7.1"
    
    # Add borders to the entire table
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Add table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    current_row = 0
    
    # Fill in education entries (content on left, graduation on right)
    for edu in education_list:
        row = table.rows[current_row]
        
        # Left cell (content)
        left_cell = row.cells[0]
        left_para = left_cell.paragraphs[0]
        
        # Degree (bold)
        degree_run = left_para.add_run(edu['degree'])
        degree_run.bold = True
        degree_run.font.size = Pt(11)
        degree_run.font.name = 'Calibri Light'
        left_para.paragraph_format.space_after = Pt(0)
        
        # Institution and location (italic)
        inst_para = left_cell.add_paragraph(f"{edu['institution']}, {edu['location']}")
        inst_para.paragraph_format.left_indent = Inches(0)
        inst_para.paragraph_format.space_before = Pt(0)
        inst_para.paragraph_format.space_after = Pt(2)
        for run in inst_para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri Light'
            run.italic = True
        
        # Add overview if present
        if edu.get('overview'):
            overview_para = left_cell.add_paragraph(f"• {edu['overview']}")
            overview_para.paragraph_format.left_indent = Inches(0.15)
            overview_para.paragraph_format.first_line_indent = Inches(-0.15)
            overview_para.paragraph_format.space_before = Pt(0)
            overview_para.paragraph_format.space_after = Pt(0)
            for run in overview_para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri Light'
        
        # Add honors/GPA if present
        if edu.get('honors') or edu.get('gpa'):
            honors_parts = []
            if edu.get('honors'):
                honors_parts.append(edu['honors'])
            if edu.get('gpa'):
                honors_parts.append(edu['gpa'])
            honors_text = '; '.join(honors_parts)
            honors_para = left_cell.add_paragraph(f"• {honors_text}")
            honors_para.paragraph_format.left_indent = Inches(0.15)
            honors_para.paragraph_format.first_line_indent = Inches(-0.15)
            honors_para.paragraph_format.space_before = Pt(0)
            honors_para.paragraph_format.space_after = Pt(0)
            for run in honors_para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri Light'
        
        # Right cell (graduation)
        right_cell = row.cells[1]
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if edu.get('graduation'):
            date_run = right_para.add_run(edu['graduation'])
            date_run.font.size = Pt(10)
            date_run.font.name = 'Calibri Light'
            date_run.italic = False
        right_para.paragraph_format.space_after = Pt(0)
        
        current_row += 1
    
    # Fill in certification entries (two per row in left column, side by side)
    cert_idx = 0
    for row_idx in range(cert_rows):
        row = table.rows[current_row + row_idx]
        
        # Fill left cell with two certifications in tabular format
        left_cell = row.cells[0]
        left_para = left_cell.paragraphs[0]
        
        # Set tab stop for alignment
        tab_stop = left_para.paragraph_format.tab_stops.add_tab_stop(Inches(3.0), WD_TAB_ALIGNMENT.LEFT)
        
        # Set hanging indent for proper wrapping
        left_para.paragraph_format.left_indent = Inches(0.15)
        left_para.paragraph_format.first_line_indent = Inches(-0.15)
        
        # First certification
        if cert_idx < len(certifications):
            cert = certifications[cert_idx]
            left_para.add_run('- ')
            name_run = left_para.add_run(f"{cert['name']}")
            name_run.bold = True
            name_run.font.size = Pt(10)
            name_run.font.name = 'Calibri Light'
            left_para.add_run(': ')
            issuer_run = left_para.add_run(cert['issuer'])
            issuer_run.font.size = Pt(10)
            issuer_run.font.name = 'Calibri Light'
            cert_idx += 1
        
        # Tab to second column
        left_para.add_run('\t')
        
        # Second certification
        if cert_idx < len(certifications):
            cert = certifications[cert_idx]
            left_para.add_run('- ')
            name_run = left_para.add_run(f"{cert['name']}")
            name_run.bold = True
            name_run.font.size = Pt(10)
            name_run.font.name = 'Calibri Light'
            left_para.add_run(': ')
            issuer_run = left_para.add_run(cert['issuer'])
            issuer_run.font.size = Pt(10)
            issuer_run.font.name = 'Calibri Light'
            cert_idx += 1
        
        left_para.paragraph_format.space_after = Pt(2)
        
        # Leave right cell empty
        right_cell = row.cells[1]
        right_para = right_cell.paragraphs[0]
        right_para.paragraph_format.space_after = Pt(2)
    
    # Add spacing after table
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_after = Pt(4)


def add_education_section_table(doc: Document, education_list: List[Dict]) -> None:
    """
    Add all education entries in a single table with shared outer border.
    Each education entry has content on left, graduation year on right.
    
    Args:
        doc: Document object
        education_list: List of education dictionaries
    """
    if not education_list:
        return
    
    # Create table with number of rows = number of education entries
    table = doc.add_table(rows=len(education_list), cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Calculate full width (page width - margins)
    full_width = Inches(7.1)
    
    # Set column widths with 75/25 ratio
    table.columns[0].width = Inches(5.325)  # 75% of 7.1"
    table.columns[1].width = Inches(1.775)  # 25% of 7.1"
    
    # Add borders to the entire table
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Add table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Fill in each row with education data
    for idx, edu in enumerate(education_list):
        row = table.rows[idx]
        
        # Left cell (content)
        left_cell = row.cells[0]
        left_para = left_cell.paragraphs[0]
        
        # Degree (bold)
        degree_run = left_para.add_run(edu['degree'])
        degree_run.bold = True
        degree_run.font.size = Pt(11)
        degree_run.font.name = 'Calibri Light'
        left_para.paragraph_format.space_after = Pt(0)
        
        # Institution and location (italic)
        inst_para = left_cell.add_paragraph(f"{edu['institution']}, {edu['location']}")
        inst_para.paragraph_format.left_indent = Inches(0)
        inst_para.paragraph_format.space_before = Pt(0)
        inst_para.paragraph_format.space_after = Pt(2)
        for run in inst_para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri Light'
            run.italic = True
        
        # Add overview if present
        if edu.get('overview'):
            overview_para = left_cell.add_paragraph(f"• {edu['overview']}")
            overview_para.paragraph_format.left_indent = Inches(0.15)
            overview_para.paragraph_format.first_line_indent = Inches(-0.15)
            overview_para.paragraph_format.space_before = Pt(0)
            overview_para.paragraph_format.space_after = Pt(0)
            for run in overview_para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri Light'
        
        # Add honors/GPA if present
        if edu.get('honors') or edu.get('gpa'):
            honors_parts = []
            if edu.get('honors'):
                honors_parts.append(edu['honors'])
            if edu.get('gpa'):
                honors_parts.append(edu['gpa'])
            honors_text = '; '.join(honors_parts)
            honors_para = left_cell.add_paragraph(f"• {honors_text}")
            honors_para.paragraph_format.left_indent = Inches(0.15)
            honors_para.paragraph_format.first_line_indent = Inches(-0.15)
            honors_para.paragraph_format.space_before = Pt(0)
            honors_para.paragraph_format.space_after = Pt(0)
            for run in honors_para.runs:
                run.font.size = Pt(10)
                run.font.name = 'Calibri Light'
        
        # Right cell (graduation)
        right_cell = row.cells[1]
        right_para = right_cell.paragraphs[0]
        right_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if edu.get('graduation'):
            date_run = right_para.add_run(edu['graduation'])
            date_run.font.size = Pt(10)
            date_run.font.name = 'Calibri Light'
            date_run.italic = False
        right_para.paragraph_format.space_after = Pt(0)
    
    # Add spacing after table
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_after = Pt(4)


def add_certification_section_table(doc: Document, certifications: List[Dict]) -> None:
    """
    Add all certification entries in a single table with shared outer border.
    Two certifications per row (2 columns).
    
    Args:
        doc: Document object
        certifications: List of certification dictionaries with keys: name, issuer
    """
    if not certifications:
        return
    
    # Calculate number of rows needed (2 certifications per row)
    num_rows = (len(certifications) + 1) // 2
    
    # Create table with 2 columns
    table = doc.add_table(rows=num_rows, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Calculate full width (page width - margins)
    full_width = Inches(7.1)
    
    # Set column widths (equal split)
    table.columns[0].width = Inches(3.55)  # 50% of 7.1"
    table.columns[1].width = Inches(3.55)  # 50% of 7.1"
    
    # Add borders to the entire table
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Add table borders
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '6')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)
    
    # Fill in cells with certification data
    cert_idx = 0
    for row_idx in range(num_rows):
        row = table.rows[row_idx]
        
        # Fill left cell
        if cert_idx < len(certifications):
            cert = certifications[cert_idx]
            left_cell = row.cells[0]
            left_para = left_cell.paragraphs[0]
            left_para.add_run('- ')
            name_run = left_para.add_run(f"{cert['name']}")
            name_run.bold = True
            name_run.font.size = Pt(10)
            name_run.font.name = 'Calibri Light'
            left_para.add_run(': ')
            issuer_run = left_para.add_run(cert['issuer'])
            issuer_run.font.size = Pt(10)
            issuer_run.font.name = 'Calibri Light'
            left_para.paragraph_format.space_after = Pt(2)
            cert_idx += 1
        
        # Fill right cell
        if cert_idx < len(certifications):
            cert = certifications[cert_idx]
            right_cell = row.cells[1]
            right_para = right_cell.paragraphs[0]
            right_para.add_run('- ')
            name_run = right_para.add_run(f"{cert['name']}")
            name_run.bold = True
            name_run.font.size = Pt(10)
            name_run.font.name = 'Calibri Light'
            right_para.add_run(': ')
            issuer_run = right_para.add_run(cert['issuer'])
            issuer_run.font.size = Pt(10)
            issuer_run.font.name = 'Calibri Light'
            right_para.paragraph_format.space_after = Pt(2)
            cert_idx += 1
    
    # Add spacing after table
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_after = Pt(4)


def add_education_entry(doc: Document,
                       degree: str,
                       institution: str,
                       location: str,
                       graduation: str,
                       gpa: Optional[str] = None,
                       honors: Optional[str] = None,
                       overview: Optional[str] = None,
                       courses_overview: Optional[str] = None,
                       courses: Optional[List[str]] = None) -> None:
    """
    Add a single education entry.
    
    Args:
        doc: Document object
        degree: Degree name
        institution: Institution name
        location: Location
        graduation: Graduation year or "In Progress"
        gpa: GPA string
        honors: Honors string
        overview: Program overview
        courses_overview: Courses overview
        courses: List of courses
    """
    # Create table for degree row (degree on left, graduation on right)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Calculate full width (page width - margins)
    # Page width is 8.5", left margin 0.7", right margin 0.7" = 7.1" available
    full_width = Inches(7.1)
    
    # Set column widths with 75/25 ratio
    table.columns[0].width = Inches(5.325)  # 75% of 7.1"
    table.columns[1].width = Inches(1.775)  # 25% of 7.1"
    
    # Remove borders
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": 0}, bottom={"sz": 0}, start={"sz": 0}, end={"sz": 0})
    
    # Degree cell (left)
    degree_cell = table.rows[0].cells[0]
    degree_para = degree_cell.paragraphs[0]
    degree_run = degree_para.add_run(f"{degree}")
    degree_run.bold = True
    degree_run.font.size = Pt(11)
    degree_run.font.name = 'Calibri Light'
    degree_para.paragraph_format.space_after = Pt(0)
    
    # Graduation cell (right)
    grad_cell = table.rows[0].cells[1]
    grad_para = grad_cell.paragraphs[0]
    grad_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    grad_run = grad_para.add_run(graduation)
    grad_run.font.size = Pt(10)
    grad_run.font.name = 'Calibri Light'
    grad_run.italic = True
    grad_para.paragraph_format.space_after = Pt(0)
    
    # Institution and location
    inst_para = doc.add_paragraph()
    inst_run = inst_para.add_run(f"{institution}, {location}")
    inst_run.font.size = Pt(10)
    inst_run.font.name = 'Calibri Light'
    inst_run.italic = True
    inst_para.paragraph_format.space_after = Pt(2)
    
    # Add overview if present
    if overview:
        overview_para = doc.add_paragraph(f"• {overview}")
        overview_para.paragraph_format.left_indent = Inches(0.25)
        overview_para.paragraph_format.first_line_indent = Inches(-0.25)
        overview_para.paragraph_format.space_after = Pt(2)
        for run in overview_para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri Light'
    
    # Add honors/GPA if present
    if honors or gpa:
        honors_parts = []
        if honors:
            honors_parts.append(honors)
        if gpa:
            honors_parts.append(gpa)
        honors_text = '; '.join(honors_parts)
        honors_para = doc.add_paragraph(f"• {honors_text}")
        honors_para.paragraph_format.left_indent = Inches(0.25)
        honors_para.paragraph_format.first_line_indent = Inches(-0.25)
        honors_para.paragraph_format.space_after = Pt(2)
        for run in honors_para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri Light'
    
    # Add courses overview if present
    if courses_overview:
        courses_para = doc.add_paragraph(f"• {courses_overview}")
        courses_para.paragraph_format.left_indent = Inches(0.25)
        courses_para.paragraph_format.first_line_indent = Inches(-0.25)
        courses_para.paragraph_format.space_after = Pt(2)
        for run in courses_para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri Light'
    
    # Add courses list if present
    if courses and len(courses) > 0:
        courses_text = ', '.join(courses)
        courses_para = doc.add_paragraph(f"• Courses: {courses_text}")
        courses_para.paragraph_format.left_indent = Inches(0.25)
        courses_para.paragraph_format.first_line_indent = Inches(-0.25)
        courses_para.paragraph_format.space_after = Pt(2)
        for run in courses_para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri Light'
    
    # Add spacing after entry
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_after = Pt(4)


def add_certification_entry(doc: Document,
                           name: str,
                           issuer: str,
                           date: str) -> None:
    """
    Add a single certification entry.
    
    Args:
        doc: Document object
        name: Certification name
        issuer: Issuing organization
        date: Issue date/year
    """
    # Create table for cert row (cert on left, date on right)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Calculate full width (page width - margins)
    # Page width is 8.5", left margin 0.7", right margin 0.7" = 7.1" available
    full_width = Inches(7.1)
    
    # Set column widths with 75/25 ratio
    table.columns[0].width = Inches(5.325)  # 75% of 7.1"
    table.columns[1].width = Inches(1.775)  # 25% of 7.1"
    
    # Remove borders
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": 0}, bottom={"sz": 0}, start={"sz": 0}, end={"sz": 0})
    
    # Cert cell (left)
    cert_cell = table.rows[0].cells[0]
    cert_para = cert_cell.paragraphs[0]
    cert_run = cert_para.add_run(f"{name}: {issuer}")
    cert_run.font.size = Pt(10)
    cert_run.font.name = 'Calibri Light'
    cert_para.paragraph_format.space_after = Pt(2)
    
    # Date cell (right)
    date_cell = table.rows[0].cells[1]
    date_para = date_cell.paragraphs[0]
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(date)
    date_run.font.size = Pt(10)
    date_run.font.name = 'Calibri Light'
    date_run.italic = True
    date_para.paragraph_format.space_after = Pt(2)


def add_volunteer_entry(doc: Document,
                       role: str,
                       organization: str,
                       location: str,
                       date_range: str,
                       description: Optional[str] = None) -> None:
    """
    Add a single volunteer experience entry in tabular format.
    
    Args:
        doc: Document object
        role: Volunteer role/position
        organization: Organization name
        location: Location
        date_range: Date range (e.g., "January 2020 - Present")
        description: Optional description of volunteer work
    """
    # Create table for title row (role on left, date on right)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Calculate full width (page width - margins)
    # Page width is 8.5", left margin 0.7", right margin 0.7" = 7.1" available
    full_width = Inches(7.1)
    
    # Set column widths with 75/25 ratio
    table.columns[0].width = Inches(5.325)  # 75% of 7.1"
    table.columns[1].width = Inches(1.775)  # 25% of 7.1"
    
    # Add vertical border between columns only
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            borders = tcPr.first_child_found_in("w:tcBorders")
            if borders is None:
                borders = OxmlElement('w:tcBorders')
                tcPr.append(borders)
            
            # Add vertical divider between cells
            if idx == 0:
                # Right border on left cell
                right = OxmlElement('w:right')
                right.set(qn('w:val'), 'single')
                right.set(qn('w:sz'), '6')
                right.set(qn('w:space'), '0')
                right.set(qn('w:color'), '000000')
                borders.append(right)
            else:
                # Left border on right cell
                left = OxmlElement('w:left')
                left.set(qn('w:val'), 'single')
                left.set(qn('w:sz'), '6')
                left.set(qn('w:space'), '0')
                left.set(qn('w:color'), '000000')
                borders.append(left)
            
            # Remove other borders
            for edge in ['top', 'bottom']:
                edge_elem = OxmlElement(f'w:{edge}')
                edge_elem.set(qn('w:val'), 'none')
                edge_elem.set(qn('w:sz'), '0')
                borders.append(edge_elem)
    
    # Role cell (left)
    role_cell = table.rows[0].cells[0]
    role_para = role_cell.paragraphs[0]
    role_run = role_para.add_run(role)
    role_run.bold = True
    role_run.font.size = Pt(11)
    role_run.font.name = 'Calibri Light'
    role_para.paragraph_format.space_after = Pt(0)
    
    # Add organization and location on the same line
    role_para.add_run(', ')
    org_run = role_para.add_run(organization)
    org_run.font.italic = True
    org_run.font.size = Pt(10)
    org_run.font.name = 'Calibri Light'
    org_run.bold = False
    
    if location and location != organization:
        role_para.add_run(', ')
        loc_run = role_para.add_run(location)
        loc_run.font.italic = True
        loc_run.font.size = Pt(10)
        loc_run.font.name = 'Calibri Light'
        loc_run.bold = False
    
    # Add description inside the left cell if present
    if description:
        desc_para = role_cell.add_paragraph(f'• {description}')
        desc_para.paragraph_format.left_indent = Inches(0.15)
        desc_para.paragraph_format.first_line_indent = Inches(-0.15)
        desc_para.paragraph_format.space_before = Pt(0)
        desc_para.paragraph_format.space_after = Pt(0)
        desc_para.paragraph_format.line_spacing = 1.0
        
        for run in desc_para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Calibri Light'
    
    # Date cell (right)
    date_cell = table.rows[0].cells[1]
    date_para = date_cell.paragraphs[0]
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(date_range)
    date_run.font.size = Pt(10)
    date_run.font.name = 'Calibri Light'
    date_run.italic = False
    date_para.paragraph_format.space_after = Pt(0)
    
    # Add spacing after table
    spacing_para = doc.add_paragraph()
    spacing_para.paragraph_format.space_after = Pt(2)


def add_skills_section(doc: Document, skills_text: str) -> None:
    """
    Add skills section as comma-separated text.
    
    Args:
        doc: Document object
        skills_text: Comma-separated skills string
    """
    add_heading_with_line(doc, "SKILLS")
    
    para = doc.add_paragraph(skills_text)
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_after = Pt(8)
    
    for run in para.runs:
        run.font.size = Pt(10)
        run.font.name = 'Calibri Light'


def set_cell_border(cell, **kwargs):
    """
    Set cell border
    
    Args:
        cell: Table cell
        **kwargs: Border settings (top, bottom, start, end)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Check for tag existance, if none found, then create one
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    # List of valid border types
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            
            # Check for tag existance, if none found, then create one
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            
            # Set border attributes
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def set_repeat_table_header(row):
    """
    Set table row to repeat as header on each page.
    
    Args:
        row: Table row
    """
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), "true")
    trPr.append(tblHeader)
    return row
