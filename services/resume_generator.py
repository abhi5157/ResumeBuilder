"""
DOCX Resume Generator using python-docx.
Creates professional resumes with precise layout control.
"""

from pathlib import Path
from typing import Optional, List, Dict, Any
from docx import Document

from models import ResumeProfile
from services.docx_utils import (
    set_page_margins,
    add_contact_header,
    add_branch_title,
    add_heading_with_line,
    add_summary_section,
    add_experience_entry,
    add_experience_section_table,
    add_education_entry,
    add_education_section_table,
    add_education_certification_combined_table,
    add_certification_entry,
    add_certification_section_table,
    add_volunteer_entry,
    add_skills_section
)
from utils.logging_utils import get_logger

logger = get_logger(__name__)


class DocxResumeGenerator:
    """
    Generate professional DOCX resumes programmatically.
    Uses python-docx for precise layout control.
    """
    
    def __init__(self):
        """Initialize the resume generator."""
        pass
    
    def generate(self, profile: ResumeProfile, output_path: Path) -> Path:
        """
        Generate a resume DOCX file from profile.
        
        Args:
            profile: Resume profile data
            output_path: Path to save the DOCX file
        
        Returns:
            Path to generated DOCX file
        """
        logger.info(f"Generating resume for {profile.contact.full_name}")
        
        # Create new document
        doc = Document()
        
        # Set page margins
        set_page_margins(doc, top=0.5, bottom=0.5, left=0.7, right=0.7)
        
        # Add contact header
        self._add_contact_section(doc, profile)
        
        # Add branch title if military background
        self._add_branch_section(doc, profile)
        
        # Add professional summary
        if profile.summary:
            add_summary_section(doc, profile.summary)
        
        # Add work experience
        if profile.work_history:
            self._add_experience_section(doc, profile)
        
        # Add education and certifications (combined section)
        if profile.education or profile.certifications:
            self._add_education_certifications_section(doc, profile)
        
        # Add skills
        if hasattr(profile, 'core_skills') or hasattr(profile, 'mos_translated_skills'):
            self._add_skills_section(doc, profile)
        
        # Add additional sections if present (volunteer, awards, etc.)
        self._add_additional_sections(doc, profile)
        
        # Save document
        doc.save(output_path)
        logger.info(f"Resume saved to {output_path}")
        
        return output_path
    
    def _add_contact_section(self, doc: Document, profile: ResumeProfile) -> None:
        """Add contact information header."""
        location = self._format_location(profile.contact.city, profile.contact.state)
        
        add_contact_header(
            doc,
            name=profile.contact.full_name,
            email=profile.contact.email,
            phone=profile.contact.phone,
            location=location,
            linkedin=profile.contact.linkedin,
            clearance=profile.contact.security_clearance if profile.contact.security_clearance != 'None' else None
        )
    
    def _add_branch_section(self, doc: Document, profile: ResumeProfile) -> None:
        """Add military branch title section."""
        branch = profile.contact.branch if hasattr(profile.contact, 'branch') and profile.contact.branch else None
        mos_code = profile.mos.code if profile.mos else ''
        mos_title = profile.mos.title if profile.mos else ''
        
        # Format branch name with "UNITED STATES" prefix
        branch_map = {
            'Army': 'UNITED STATES ARMY',
            'Navy': 'UNITED STATES NAVY',
            'Marines': 'UNITED STATES MARINE CORPS',
            'Marine Corps': 'UNITED STATES MARINE CORPS',
            'Air Force': 'UNITED STATES AIR FORCE',
            'Space Force': 'UNITED STATES SPACE FORCE',
            'Coast Guard': 'UNITED STATES COAST GUARD'
        }
        
        if branch and mos_title:
            full_branch = branch_map.get(branch, branch.upper())
            branch_title = f"{full_branch} - {mos_title.upper()}"
            add_branch_title(doc, branch_title)
        elif branch:
            full_branch = branch_map.get(branch, branch.upper())
            branch_title = f"{full_branch} VETERAN"
            add_branch_title(doc, branch_title)
    
    def _add_experience_section(self, doc: Document, profile: ResumeProfile) -> None:
        """Add work experience section."""
        add_heading_with_line(doc, "PROFESSIONAL EXPERIENCE")
        
        # Prepare all experience entries
        experiences = []
        for exp in profile.work_history:
            # Use AI-generated bullets if available, otherwise use original bullets
            bullets = exp.ai_generated_bullets if exp.ai_generated_bullets else exp.achievements
            
            experiences.append({
                'title': exp.job_title,
                'subtitle': exp.employer,
                'location': exp.location or '',
                'date_range': exp.date_range,
                'bullets': bullets
            })
        
        # Add all experiences in a single table with shared outer border
        add_experience_section_table(doc, experiences)
    
    def _add_education_certifications_section(self, doc: Document, profile: ResumeProfile) -> None:
        """Add combined education and certifications section."""
        add_heading_with_line(doc, "EDUCATION & CERTIFICATIONS")
        
        # Prepare all education entries
        education_list = []
        for edu in profile.education:
            # Use date_range if available, otherwise use graduation_year
            grad_info = ""
            if hasattr(edu, 'date_range') and edu.date_range:
                grad_info = edu.date_range
            elif edu.in_progress:
                grad_info = "In Progress"
            elif edu.graduation_year:
                grad_info = str(edu.graduation_year)
            
            # Get location for education
            edu_location = ""
            if hasattr(edu, 'city') and hasattr(edu, 'state'):
                edu_location = self._format_location(edu.city, edu.state)
            elif hasattr(edu, 'location'):
                edu_location = edu.location
            
            # Format GPA and honors
            gpa_str = f"GPA: {edu.gpa:.2f}" if edu.gpa else None
            honors_str = ', '.join(edu.honors) if edu.honors else None
            
            education_list.append({
                'degree': edu.degree,
                'institution': edu.institution,
                'location': edu_location,
                'graduation': grad_info,
                'gpa': gpa_str,
                'honors': honors_str,
                'overview': edu.overview,
                'courses_overview': getattr(edu, 'courses_overview', None),
                'courses': getattr(edu, 'courses', None)
            })
        
        # Prepare all certification entries (no date field)
        certifications = []
        for cert in profile.certifications:
            certifications.append({
                'name': cert.name,
                'issuer': cert.issuer
            })
        
        # Add both education and certifications in a single table with shared outer border
        add_education_certification_combined_table(doc, education_list, certifications)
    
    def _add_certifications_section(self, doc: Document, profile: ResumeProfile) -> None:
        """Add certifications section."""
        add_heading_with_line(doc, "CERTIFICATIONS")
        
        # Prepare all certification entries
        certifications = []
        for cert in profile.certifications:
            certifications.append({
                'name': cert.name,
                'issuer': cert.issuer
            })
        
        # Add certifications in a table (two per row)
        add_certification_section_table(doc, certifications)
    
    def _add_skills_section(self, doc: Document, profile: ResumeProfile) -> None:
        """Add skills section."""
        skills_list = []
        
        # Add MOS-translated skills first (most relevant)
        if hasattr(profile, 'mos_translated_skills') and profile.mos_translated_skills:
            skills_list.extend(profile.mos_translated_skills)
        
        # Add core skills
        if hasattr(profile, 'core_skills') and profile.core_skills:
            for skill in profile.core_skills:
                if skill not in skills_list:
                    skills_list.append(skill)
        
        # Add tools and technologies
        if hasattr(profile, 'tools_technologies') and profile.tools_technologies:
            for skill in profile.tools_technologies:
                if skill not in skills_list:
                    skills_list.append(skill)
        
        # Add target keywords for ATS
        if hasattr(profile, 'target_keywords') and profile.target_keywords:
            for skill in profile.target_keywords:
                if skill not in skills_list:
                    skills_list.append(skill)
        
        # Add MOS civilian skills if available
        if profile.mos and hasattr(profile.mos, 'civilian_skills') and profile.mos.civilian_skills:
            for skill in profile.mos.civilian_skills:
                if skill not in skills_list:
                    skills_list.append(skill)
        
        # Format skills as comma-separated text
        if skills_list:
            skills_text = ', '.join(skills_list)
            add_skills_section(doc, skills_text)
    
    def _add_additional_sections(self, doc: Document, profile: ResumeProfile) -> None:
        """Add additional optional sections (awards, volunteer, etc.)."""
        if not profile.additional_info:
            return
        
        # Add awards if present
        if profile.additional_info.awards and len(profile.additional_info.awards) > 0:
            add_heading_with_line(doc, "AWARDS & HONORS")
            for award in profile.additional_info.awards:
                para = doc.add_paragraph(f"• {award}")
                para.paragraph_format.left_indent = Inches(0.25)
                para.paragraph_format.first_line_indent = Inches(-0.25)
                para.paragraph_format.space_after = Pt(2)
                for run in para.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri Light'
        
        # Add volunteer experience if present
        if profile.additional_info.volunteer and len(profile.additional_info.volunteer) > 0:
            add_heading_with_line(doc, "VOLUNTEER EXPERIENCE")
            for vol in profile.additional_info.volunteer:
                if isinstance(vol, str):
                    # Handle simple string format
                    para = doc.add_paragraph(f"• {vol}")
                    para.paragraph_format.left_indent = Inches(0.25)
                    para.paragraph_format.first_line_indent = Inches(-0.25)
                    para.paragraph_format.space_after = Pt(2)
                    for run in para.runs:
                        run.font.size = Pt(10)
                        run.font.name = 'Calibri Light'
                elif isinstance(vol, dict):
                    # Handle dict format from JSON
                    role = vol.get('role', 'Volunteer')
                    organization = vol.get('organization', '')
                    location = vol.get('location', '')
                    date_range = vol.get('date_range', '')
                    description = vol.get('description', None)
                    
                    add_volunteer_entry(
                        doc,
                        role=role,
                        organization=organization,
                        location=location,
                        date_range=date_range,
                        description=description
                    )
                else:
                    # Handle VolunteerExperience object
                    role = getattr(vol, 'role', 'Volunteer')
                    organization = getattr(vol, 'organization', '')
                    location = getattr(vol, 'location', '')
                    date_range = getattr(vol, 'date_range', '')
                    description = getattr(vol, 'description', None)
                    
                    add_volunteer_entry(
                        doc,
                        role=role,
                        organization=organization,
                        location=location,
                        date_range=date_range,
                        description=description
                    )
    
    def _format_location(self, city: Optional[str], state: Optional[str]) -> str:
        """Format city and state for display."""
        if city and state:
            return f"{city}, {state}"
        elif city:
            return city
        elif state:
            return state
        return ''


# Import needed for docx_utils
from docx.shared import Inches, Pt
